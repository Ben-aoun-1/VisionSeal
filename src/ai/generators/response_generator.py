"""
Enhanced Response Generation for VisionSeal RAG System
Generates professional tender responses with Topaza branding
"""
import os
import json
import re
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path

# Document generation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# AI and utilities
import tiktoken
from openai import OpenAI

from core.config.settings import settings
from core.logging.setup import get_logger
from core.exceptions.handlers import AIProcessingException

logger = get_logger("response_generator")


class ResponseGenerator:
    """Enhanced response generator with enterprise features and advanced prompt engineering"""
    
    def __init__(
        self,
        openai_api_key: Optional[str] = None,
        model: str = "gpt-4-turbo",
        max_context_tokens: int = 3000,
        max_response_tokens: int = 4000
    ):
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        if not self.openai_api_key:
            raise AIProcessingException("OpenAI API key not provided")
            
        self.openai_client = OpenAI(api_key=self.openai_api_key)
        self.model = model
        self.max_context_tokens = max_context_tokens
        self.max_response_tokens = max_response_tokens
        
        # Initialize tokenizer for precise token counting
        self.tokenizer = tiktoken.encoding_for_model("gpt-4")
        
        # Length specifications for different report types
        self.length_specs = {
            "brief": {
                "target_words": 800,
                "max_tokens": 1200,
                "sections": 4
            },
            "detailed": {
                "target_words": 2500,
                "max_tokens": 3500,
                "sections": 8
            },
            "comprehensive": {
                "target_words": 5000,
                "max_tokens": 7000,
                "sections": 12
            }
        }

    async def generate_tender_summary(self, tender_text: str) -> str:
        """Generate structured summary of tender document"""
        try:
            # Limit text to avoid token overflow
            limited_text = tender_text[:15000] if len(tender_text) > 15000 else tender_text
            
            summary_prompt = f"""
            Vous êtes un expert en analyse d'appels d'offres. Générez un résumé structuré qui met en avant:
            
            ## Analyse de l'Appel d'Offres
            
            ### Exigences Clés
            - [Listez les exigences principales]
            
            ### Normes et Standards
            - [Normes techniques requises]
            
            ### Délais et Échéances
            - [Dates importantes]
            
            ### Livrables Attendus
            - [Productions à fournir]
            
            ### Critères d'Évaluation
            - [Critères de sélection]
            
            ### Informations Stratégiques
            - [Autres éléments cruciaux]
            
            Texte de l'AO:
            {limited_text}
            
            Fournissez le résumé selon la structure exacte ci-dessus.
            """
            
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Expert en analyse d'appels d'offres pour l'entreprise Topaza"},
                    {"role": "user", "content": summary_prompt}
                ],
                temperature=0.3,
                max_tokens=3500
            )
            
            summary = response.choices[0].message.content.strip()
            logger.info("Tender summary generated successfully")
            return summary
            
        except Exception as e:
            logger.error(f"Tender summary generation failed: {str(e)}")
            raise AIProcessingException(f"Tender summary generation failed: {str(e)}")
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using the model's tokenizer"""
        try:
            return len(self.tokenizer.encode(text))
        except Exception:
            # Fallback to character-based estimation
            return len(text) // 4
    
    def _get_tone_instructions(self, tone: str) -> str:
        """Get tone-specific instructions for response generation"""
        tone_map = {
            "professional": {
                "style": "Adoptez un ton professionnel, formel et respectueux. Utilisez un vocabulaire technique approprié et une structure claire.",
                "phrases": ["nous proposons", "notre expertise", "conformément aux exigences", "dans le respect des délais"]
            },
            "technical": {
                "style": "Utilisez un langage technique précis avec des détails spécifiques. Mettez l'accent sur les aspects techniques et méthodologiques.",
                "phrases": ["selon les spécifications", "architecture technique", "méthodologie éprouvée", "standards industriels"]
            },
            "persuasive": {
                "style": "Adoptez un ton convaincant qui met en valeur nos atouts. Utilisez des arguments persuasifs tout en restant factuel.",
                "phrases": ["solution optimale", "avantage concurrentiel", "valeur ajoutée", "retour sur investissement"]
            }
        }
        
        tone_config = tone_map.get(tone, tone_map["professional"])
        return f"""
        STYLE ET TON : {tone_config['style']}
        
        EXPRESSIONS RECOMMANDÉES : {', '.join(tone_config['phrases'])}
        
        APPROCHE : Adaptez le niveau de détail et le vocabulaire selon ce ton tout au long de la réponse.
        """
    
    def _get_report_type_structure(self, report_type: str, length: str) -> Dict[str, Any]:
        """Get structure and sections based on report type and length"""
        structures = {
            "proposal": {
                "brief": {
                    "sections": [
                        "Résumé Exécutif",
                        "Approche Méthodologique",
                        "Équipe et Expertise",
                        "Planning et Budget"
                    ],
                    "focus": "Présentation concise de notre proposition technique et commerciale"
                },
                "detailed": {
                    "sections": [
                        "Résumé Exécutif",
                        "Compréhension du Besoin",
                        "Approche Méthodologique Détaillée",
                        "Plan de Gestion des Risques",
                        "Équipe Projet et Expertise",
                        "Planning Détaillé",
                        "Budget et Justifications",
                        "Garanties et Suivi"
                    ],
                    "focus": "Proposition complète avec détails techniques et méthodologiques"
                },
                "comprehensive": {
                    "sections": [
                        "Résumé Exécutif",
                        "Analyse du Contexte et des Enjeux",
                        "Compréhension Approfondie du Besoin",
                        "Méthodologie et Approche Technique",
                        "Architecture de la Solution",
                        "Plan de Gestion des Risques et Mitigation",
                        "Organisation de l'Équipe Projet",
                        "Planning Détaillé par Phase",
                        "Budget Détaillé et Justifications",
                        "Garanties de Qualité et Suivi",
                        "Livrables et Documentation",
                        "Valeur Ajoutée et Innovation"
                    ],
                    "focus": "Proposition exhaustive avec tous les aspects techniques, organisationnels et commerciaux"
                }
            },
            "analysis": {
                "brief": {
                    "sections": [
                        "Synthèse de l'Appel d'Offres",
                        "Analyse des Opportunités",
                        "Évaluation de la Faisabilité",
                        "Recommandations Stratégiques"
                    ],
                    "focus": "Analyse rapide des opportunités et recommandations"
                },
                "detailed": {
                    "sections": [
                        "Synthèse Exécutive",
                        "Analyse Détaillée des Exigences",
                        "Évaluation du Marché et de la Concurrence",
                        "Analyse de Faisabilité Technique",
                        "Évaluation des Risques et Opportunités",
                        "Analyse Financière Préliminaire",
                        "Recommandations Stratégiques",
                        "Plan d'Action Proposé"
                    ],
                    "focus": "Analyse complète avec évaluation technique et commerciale"
                },
                "comprehensive": {
                    "sections": [
                        "Synthèse Exécutive",
                        "Analyse Contextuelle et Stratégique",
                        "Décryptage Détaillé des Exigences",
                        "Étude de Marché et Positionnement Concurrentiel",
                        "Analyse de Faisabilité Multi-dimensionnelle",
                        "Évaluation Approfondie des Risques",
                        "Analyse Financière et ROI",
                        "Étude d'Impact et de Conformité",
                        "Scénarios et Alternatives",
                        "Recommandations Stratégiques Détaillées",
                        "Feuille de Route et Planning",
                        "Conclusions et Perspectives"
                    ],
                    "focus": "Analyse exhaustive sous tous les angles avec recommandations détaillées"
                }
            },
            "summary": {
                "brief": {
                    "sections": [
                        "Informations Clés",
                        "Exigences Principales",
                        "Délais et Contraintes",
                        "Points d'Attention"
                    ],
                    "focus": "Synthèse rapide des éléments essentiels"
                },
                "detailed": {
                    "sections": [
                        "Vue d'Ensemble",
                        "Objectifs et Enjeux",
                        "Exigences Techniques Détaillées",
                        "Contraintes et Délais",
                        "Critères d'Évaluation",
                        "Livrables Attendus",
                        "Modalités de Participation",
                        "Points Critiques à Retenir"
                    ],
                    "focus": "Synthèse structurée avec tous les éléments importants"
                },
                "comprehensive": {
                    "sections": [
                        "Synthèse Exécutive",
                        "Contexte et Enjeux Stratégiques",
                        "Objectifs Détaillés et Périmètre",
                        "Spécifications Techniques Complètes",
                        "Exigences Qualité et Conformité",
                        "Planning et Jalons Critiques",
                        "Budget et Modalités Financières",
                        "Critères d'Évaluation et Pondération",
                        "Procédure de Soumission",
                        "Risques et Mitigation",
                        "Opportunités Identifiées",
                        "Recommandations pour la Réponse"
                    ],
                    "focus": "Synthèse exhaustive avec analyse approfondie de tous les aspects"
                }
            }
        }
        
        return structures.get(report_type, structures["proposal"]).get(length, structures["proposal"]["detailed"])
    
    async def generate_intelligent_report(
        self,
        tender_data: Dict[str, Any],
        report_type: str = "proposal",
        tone: str = "professional",
        length: str = "detailed",
        custom_instructions: Optional[str] = None,
        rag_context: Optional[List[str]] = None
    ) -> str:
        """Generate intelligent, well-structured report based on tender data and requirements"""
        try:
            logger.info(f"Generating {length} {report_type} report with {tone} tone")
            
            # Get report structure
            structure = self._get_report_type_structure(report_type, length)
            length_spec = self.length_specs[length]
            
            # Prepare context
            tender_text = self._prepare_tender_context(tender_data)
            rag_context_text = "\n".join(rag_context) if rag_context else ""
            
            # Generate report in sections for better quality and length control
            sections = []
            total_words = 0
            target_words_per_section = length_spec["target_words"] // len(structure["sections"])
            
            for i, section_title in enumerate(structure["sections"]):
                logger.info(f"Generating section {i+1}/{len(structure['sections'])}: {section_title}")
                
                section_content = await self._generate_section(
                    section_title=section_title,
                    tender_data=tender_data,
                    tender_text=tender_text,
                    rag_context=rag_context_text,
                    report_type=report_type,
                    tone=tone,
                    target_words=target_words_per_section,
                    custom_instructions=custom_instructions,
                    section_number=i+1,
                    total_sections=len(structure["sections"]),
                    structure_focus=structure["focus"]
                )
                
                sections.append(section_content)
                section_words = len(section_content.split())
                total_words += section_words
                
                logger.info(f"Section '{section_title}' generated: {section_words} words")
            
            # Combine all sections
            full_report = "\n\n".join(sections)
            
            # Add header with metadata
            header = self._generate_report_header(tender_data, report_type, tone, length)
            final_report = f"{header}\n\n{full_report}"
            
            logger.info(f"Report generation completed: {total_words} words, {len(sections)} sections")
            
            return final_report
            
        except Exception as e:
            logger.error(f"Intelligent report generation failed: {str(e)}")
            raise AIProcessingException(f"Report generation failed: {str(e)}")
    
    def _prepare_tender_context(self, tender_data: Dict[str, Any]) -> str:
        """Prepare tender context from data"""
        context_parts = []
        
        if tender_data.get("title"):
            context_parts.append(f"TITRE: {tender_data['title']}")
        
        if tender_data.get("description"):
            context_parts.append(f"DESCRIPTION: {tender_data['description']}")
        
        if tender_data.get("organization"):
            context_parts.append(f"ORGANISATION: {tender_data['organization']}")
        
        if tender_data.get("country"):
            context_parts.append(f"PAYS: {tender_data['country']}")
        
        if tender_data.get("deadline"):
            context_parts.append(f"ÉCHÉANCE: {tender_data['deadline']}")
        
        if tender_data.get("reference"):
            context_parts.append(f"RÉFÉRENCE: {tender_data['reference']}")
        
        return "\n".join(context_parts)
    
    async def _generate_section(
        self,
        section_title: str,
        tender_data: Dict[str, Any],
        tender_text: str,
        rag_context: str,
        report_type: str,
        tone: str,
        target_words: int,
        custom_instructions: Optional[str],
        section_number: int,
        total_sections: int,
        structure_focus: str
    ) -> str:
        """Generate a specific section with advanced prompt engineering"""
        
        # Build sophisticated prompt
        prompt = f"""Vous êtes un expert consultant senior spécialisé dans la rédaction de réponses aux appels d'offres pour l'entreprise TOPAZA.

CONTEXTE DE L'APPEL D'OFFRES:
{tender_text}

CONTEXTE RAG (Documents de référence):
{rag_context[:2000] if rag_context else "Aucun contexte additionnel"}

INSTRUCTIONS DE GÉNÉRATION:
- Type de rapport: {report_type.upper()}
- Section à générer: {section_title} (Section {section_number}/{total_sections})
- Objectif du rapport: {structure_focus}
- Longueur cible pour cette section: {target_words} mots environ

{self._get_tone_instructions(tone)}

INSTRUCTIONS SPÉCIFIQUES POUR CETTE SECTION:
{self._get_section_specific_instructions(section_title, report_type)}

INSTRUCTIONS PERSONNALISÉES:
{custom_instructions if custom_instructions else "Aucune instruction personnalisée"}

EXIGENCES DE QUALITÉ:
1. Utilisez EXCLUSIVEMENT les informations de l'appel d'offres et du contexte fourni
2. Adaptez le contenu spécifiquement à TOPAZA et ses capacités
3. Respectez la longueur cible de {target_words} mots (±20%)
4. Structurez avec des sous-sections numérotées
5. Utilisez des listes à puces pour la clarté
6. Incluez des détails concrets et spécifiques
7. Évitez les généralités et les formulations vagues

STRUCTURE ATTENDUE:
## {section_number}. {section_title}

[Générez le contenu de cette section en respectant toutes les exigences ci-dessus]

IMPORTANT: Générez UNIQUEMENT le contenu de cette section, pas d'introduction ni de conclusion générale."""

        try:
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system", 
                        "content": "Vous êtes un expert en rédaction d'appels d'offres avec 15 ans d'expérience. Votre spécialité est de créer des réponses techniques de haute qualité, structurées et persuasives."
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                temperature=0.3,  # Lower temperature for more focused output
                max_tokens=min(2500, target_words * 2),  # Allow room for detailed content
                presence_penalty=0.1,  # Encourage diverse content
                frequency_penalty=0.1   # Reduce repetition
            )
            
            content = response.choices[0].message.content.strip()
            
            # Validate section length and quality
            word_count = len(content.split())
            if word_count < target_words * 0.7:
                logger.warning(f"Section '{section_title}' is shorter than expected: {word_count} words")
            
            return content
            
        except Exception as e:
            logger.error(f"Section generation failed for '{section_title}': {str(e)}")
            return f"## {section_number}. {section_title}\n\nErreur lors de la génération de cette section. Veuillez réessayer."
    
    def _get_section_specific_instructions(self, section_title: str, report_type: str) -> str:
        """Get specific instructions for each section type"""
        instructions_map = {
            "Résumé Exécutif": """
            - Présentez en 2-3 paragraphes notre compréhension de l'AO et notre proposition de valeur
            - Mettez en avant nos points forts et notre différenciation
            - Mentionnez brièvement la méthodologie et les livrables clés
            - Terminez par une phrase d'engagement fort
            """,
            
            "Compréhension du Besoin": """
            - Reformulez les objectifs de l'AO avec votre interprétation
            - Identifiez les enjeux sous-jacents et les défis techniques
            - Montrez votre compréhension des contraintes et du contexte
            - Mettez en évidence les points critiques de succès
            """,
            
            "Approche Méthodologique": """
            - Décrivez votre méthodologie étape par étape
            - Justifiez chaque choix méthodologique
            - Incluez les outils, techniques et frameworks utilisés
            - Montrez comment votre approche répond aux exigences spécifiques
            """,
            
            "Équipe et Expertise": """
            - Présentez les profils clés avec leurs expertises spécifiques
            - Montrez l'adéquation entre l'équipe et les besoins de l'AO
            - Incluez les certifications et expériences pertinentes
            - Décrivez l'organisation de l'équipe et les responsabilités
            """,
            
            "Planning et Budget": """
            - Proposez un planning réaliste avec jalons clés
            - Justifiez les délais proposés
            - Pour le budget, montrez la répartition par poste
            - Mettez en avant le rapport qualité/prix
            """
        }
        
        # Use fuzzy matching for section titles
        for key, instructions in instructions_map.items():
            if any(word in section_title for word in key.split()):
                return instructions
        
        # Default instructions
        return """
        - Développez le contenu de manière structurée et détaillée
        - Utilisez des exemples concrets liés à l'appel d'offres
        - Montrez votre expertise et votre valeur ajoutée
        - Restez factuel et professionnel
        """
    
    def _generate_report_header(self, tender_data: Dict[str, Any], report_type: str, tone: str, length: str) -> str:
        """Generate report header with metadata"""
        type_labels = {
            "proposal": "Proposition Technique et Commerciale",
            "analysis": "Analyse d'Appel d'Offres", 
            "summary": "Synthèse d'Appel d'Offres"
        }
        
        length_labels = {
            "brief": "Version Synthétique",
            "detailed": "Version Détaillée",
            "comprehensive": "Version Complète"
        }
        
        tone_labels = {
            "professional": "Professionnel",
            "technical": "Technique", 
            "persuasive": "Persuasif"
        }
        
        header = f"""# {type_labels.get(report_type, 'Rapport')} - {length_labels.get(length, 'Standard')}

**Appel d'Offres :** {tender_data.get('title', 'Non spécifié')}
**Organisation :** {tender_data.get('organization', 'Non spécifiée')}
**Pays :** {tender_data.get('country', 'Non spécifié')}
**Référence :** {tender_data.get('reference', 'Non spécifiée')}
**Échéance :** {tender_data.get('deadline', 'Non spécifiée')}

**Type de rapport :** {type_labels.get(report_type)}
**Ton :** {tone_labels.get(tone)}
**Niveau de détail :** {length_labels.get(length)}

**Généré par :** TOPAZA AI Assistant
**Date de génération :** {datetime.now().strftime('%d/%m/%Y à %H:%M')}

---
"""
        return header

    async def format_context_chunks(self, chunks: List[Dict[str, Any]]) -> str:
        """Format retrieved chunks for prompt context"""
        try:
            formatted = []
            token_count = 0
            
            for chunk in chunks:
                # Build source information
                source_info = f"Source: {chunk.get('filename', 'Unknown')}\nSection: {chunk.get('section', 'Unknown')}"
                
                # Handle table content
                content = chunk.get('content', '')
                if '[TABLE_DATA]' in content:
                    table_content = self._extract_table_content(content)
                    formatted_content = f"{source_info}\nTableau:\n{table_content}"
                else:
                    # Truncate long content
                    if len(content) > 1500:
                        content = content[:1500] + "... [contenu tronqué]"
                    formatted_content = f"{source_info}\nContenu:\n{content}"
                
                # Check token limits
                new_tokens = self._count_tokens(formatted_content)
                if token_count + new_tokens > self.max_context_tokens:
                    break
                
                formatted.append(formatted_content)
                token_count += new_tokens
            
            result = "\n\n---\n\n".join(formatted)
            logger.info(f"Context formatted: {len(formatted)} chunks, {token_count} tokens")
            return result
            
        except Exception as e:
            logger.error(f"Context formatting failed: {str(e)}")
            raise AIProcessingException(f"Context formatting failed: {str(e)}")

    def _extract_table_content(self, content: str) -> str:
        """Extract and format table content from marked text"""
        try:
            # Look for table markers
            table_match = re.search(r'\[TABLE_DATA\](.*?)\[/TABLE_DATA\]', content, re.DOTALL)
            if not table_match:
                return "Tableau (format non reconnu)"
            
            table_json = table_match.group(1)
            table_data = json.loads(table_json)
            return self._format_table_for_prompt(table_data)
            
        except Exception as e:
            logger.warning(f"Table extraction failed: {str(e)}")
            return "Tableau (contenu non analysable)"

    def _format_table_for_prompt(self, table_data: Any) -> str:
        """Format table data for prompt inclusion"""
        if not table_data:
            return "Aucune donnée de tableau"
        
        try:
            rows = []
            if isinstance(table_data, list) and table_data:
                # Use first row as headers
                headers = list(table_data[0].keys()) if isinstance(table_data[0], dict) else []
                if headers:
                    rows.append(" | ".join(headers))
                    rows.append(" | ".join(["---"] * len(headers)))
                    
                    for row in table_data:
                        if isinstance(row, dict):
                            rows.append(" | ".join(str(row.get(h, "")) for h in headers))
                        else:
                            rows.append(str(row))
            elif isinstance(table_data, dict):
                for key, value in table_data.items():
                    rows.append(f"{key} | {value}")
            else:
                rows.append(str(table_data))
            
            return "\n".join(rows)
            
        except Exception as e:
            logger.warning(f"Table formatting failed: {str(e)}")
            return str(table_data)

    async def build_section_prompt(
        self,
        tender_summary: str,
        context: str,
        company_profile: str,
        section_title: str,
        section_requirements: str
    ) -> str:
        """Build sophisticated prompt for section generation"""
        try:
            # Define section-specific instructions
            section_instructions = self._get_section_instructions(section_title)
            
            # Build comprehensive prompt
            prompt = f"""
            Vous êtes un expert en réponse aux appels d'offres pour l'entreprise Topaza International.
            
            Votre mission : générer la section **{section_title}** d'une réponse professionnelle.
            
            ---
            
            **PROFIL DE TOPAZA :**
            {company_profile[:2000]}
            
            **RÉSUMÉ DE L'APPEL D'OFFRES :**
            {tender_summary}
            
            **CONTEXTE DE RÉPONSES SIMILAIRES :**
            {context}
            
            **EXIGENCES POUR CETTE SECTION :**
            {section_requirements}
            
            **INSTRUCTIONS RÉDACTIONNELLES :**
            {self._get_writing_guidelines()}
            
            **FORMATAGE :**
            {self._get_formatting_instructions()}
            
            {section_instructions}
            
            ---
            
            Générez une section professionnelle, détaillée et adaptée aux exigences spécifiques de cet appel d'offres.
            Utilisez le style Topaza et intégrez les éléments du contexte pertinents.
            """
            
            return prompt
            
        except Exception as e:
            logger.error(f"Prompt building failed: {str(e)}")
            raise AIProcessingException(f"Prompt building failed: {str(e)}")

    def _get_section_instructions(self, section_title: str) -> str:
        """Get specific instructions based on section type"""
        section_lower = section_title.lower()
        
        if "risque" in section_lower or "conclusion" in section_lower:
            return """
            **IMPORTANT - Gestion des Risques :**
            - OBLIGATOIRE : Créer un tableau avec les colonnes : FACTEURS | Probabilité(/5) | Gravité(/5) | Criticité % | Eval. | Mesure(s)
            - Calculer Criticité % = (Probabilité × Gravité × 4)%
            - Eval. : Faible (<20%), Moyen (20-40%), Élevé (>40%)
            - Minimum 5 risques identifiés
            - Format tableau Markdown obligatoire
            
            **IMPORTANT - Équipe :**
            - OBLIGATOIRE : Tableau avec colonnes : Nom et Prénom | Poste | Compétences | Expérience
            - Format tableau Markdown obligatoire
            
            **IMPORTANT - Conclusion :**
            - Intégrer les contacts Topaza :
              Email : contact@topaza.net
              Téléphone : +216 90 15 12 30
              Site web : www.topaza.net
            """
        elif "méthodologie" in section_lower or "livrable" in section_lower:
            return """
            **IMPORTANT - Méthodologie :**
            - Développer une approche structurée en phases détaillées
            - Décrire précisément chaque phase avec activités, méthodes et outils
            - Présenter les livrables pour chaque phase
            - Inclure planning et chronogramme
            - Méthodes de suivi et d'évaluation
            - NE PAS inclure de tableaux de risques ou d'équipe dans cette section
            """
        else:
            return """
            **IMPORTANT - Introduction/Objectifs :**
            - Présentation détaillée de Topaza et ses valeurs
            - Analyse approfondie des objectifs de la mission
            - Démarche proposée avec sous-sections détaillées
            - Compréhension des enjeux et défis du projet
            """

    def _get_writing_guidelines(self) -> str:
        """Get Topaza writing style guidelines"""
        return """
        Style rédactionnel Topaza :
        
        * Rédaction claire, structurée, fluide et professionnelle
        * Titres hiérarchisés numérotés (1, 1.1, 2.1.1, ...)
        * Paragraphes cohérents avec transitions logiques
        * Listes à puces pour les éléments clés
        * Tableaux professionnels pour planning, équipes, livrables
        * Ton engageant et positif, formulation active ("nous proposons" au lieu de "nous pouvons")
        * Mise en avant des atouts distinctifs de Topaza :
          • Approche sur mesure
          • Expertise sectorielle
          • Méthodologies éprouvées
          • Suivi rigoureux
          • Résultats tangibles
        """

    def _get_formatting_instructions(self) -> str:
        """Get formatting instructions"""
        return """
        **FORMATAGE OBLIGATOIRE DES TITRES :**
        - Titre principal : ## 1. [Nom de la section]
        - Sous-titre niveau 1 : ### 1.1 [Nom du sous-titre]
        - Sous-titre niveau 2 : #### 1.1.1 [Nom du sous-titre]
        - TOUJOURS utiliser cette numérotation hiérarchique
        - TOUJOURS commencer les titres par ##, ###, ou ####
        
        **DÉVELOPPEMENT DU CONTENU :**
        - Rédigez de manière détaillée et approfondie chaque section
        - Utilisez des exemples concrets et des détails techniques
        - Structurez le contenu avec des sous-sections claires
        - Développez les aspects pratiques et opérationnels
        - Minimum 3-4 paragraphes par sous-section importante
        """

    async def generate_section(
        self,
        tender_summary: str,
        context: str,
        company_profile: str,
        section_title: str,
        section_requirements: str
    ) -> str:
        """Generate a specific section of the tender response"""
        try:
            prompt = await self.build_section_prompt(
                tender_summary, context, company_profile, 
                section_title, section_requirements
            )
            
            # Check token limits
            prompt_tokens = self._count_tokens(prompt)
            if prompt_tokens > 8000:  # Leave room for response
                logger.warning(f"Prompt too long ({prompt_tokens} tokens), truncating context")
                # Truncate context and rebuild
                truncated_context = context[:1000] + "... [contexte tronqué]"
                prompt = await self.build_section_prompt(
                    tender_summary, truncated_context, company_profile,
                    section_title, section_requirements
                )
            
            logger.info(f"Generating section: {section_title}")
            
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Expert Topaza en réponse aux appels d'offres"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=self.max_response_tokens
            )
            
            content = response.choices[0].message.content.strip()
            logger.info(f"Section generated successfully: {len(content)} characters")
            return content
            
        except Exception as e:
            logger.error(f"Section generation failed: {str(e)}")
            raise AIProcessingException(f"Section generation failed: {str(e)}")

    async def save_to_docx(
        self, 
        content: str, 
        filename: str,
        output_dir: Optional[str] = None
    ) -> str:
        """Save generated content to professionally formatted Word document"""
        try:
            if not output_dir:
                output_dir = settings.data_dir / "ai_responses"
            
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            full_path = output_path / filename
            
            # Create Word document
            doc = Document()
            
            # Add cover page
            self._add_cover_page(doc)
            doc.add_page_break()
            
            # Add table of contents placeholder
            self._add_table_of_contents(doc)
            doc.add_page_break()
            
            # Set default style
            self._set_default_style(doc)
            
            # Process content blocks
            self._process_content_blocks(doc, content)
            
            # Save document
            doc.save(str(full_path))
            
            logger.info(f"Document saved successfully: {full_path}")
            return str(full_path)
            
        except Exception as e:
            logger.error(f"Document saving failed: {str(e)}")
            raise AIProcessingException(f"Document saving failed: {str(e)}")

    def _add_cover_page(self, doc: Document):
        """Add professional cover page"""
        # Company title
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run("TOPAZA INTERNATIONAL")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Tagline
        tagline = doc.add_paragraph("Empower your impact!")
        tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Document title
        doc_title = doc.add_paragraph()
        doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = doc_title.add_run("Réponse à Appel d'Offre")
        run.bold = True
        run.font.size = Pt(16)
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"Date : {datetime.now().strftime('%d %B %Y')}")
        date_run.font.size = Pt(12)

    def _add_table_of_contents(self, doc: Document):
        """Add table of contents placeholder"""
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = toc_title.add_run("Table des matières")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add TOC field (will need manual update in Word)
        p = doc.add_paragraph()
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

    def _set_default_style(self, doc: Document):
        """Set default document style"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _process_content_blocks(self, doc: Document, content: str):
        """Process content blocks and add to document with proper formatting"""
        blocks = content.split('\n\n')
        
        for block in blocks:
            block_clean = block.strip()
            if not block_clean:
                continue
            
            # Remove markdown bold formatting
            block_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', block_clean)
            
            # Check if block contains a table
            if self._is_table_block(block_clean):
                self._add_table_to_doc(doc, block_clean)
            else:
                self._add_text_block_to_doc(doc, block_clean)

    def _is_table_block(self, block: str) -> bool:
        """Check if block contains a table"""
        lines = block.split('\n')
        table_indicators = sum(1 for line in lines if '|' in line)
        return table_indicators >= 2 and ('FACTEURS' in block or 'Nom et Prénom' in block)

    def _add_table_to_doc(self, doc: Document, block: str):
        """Add formatted table to document"""
        try:
            lines = [line.strip() for line in block.split('\n') if '|' in line]
            if len(lines) < 2:
                doc.add_paragraph(block)
                return
            
            # Extract headers and data
            headers = [h.strip() for h in lines[0].split('|')]
            data_rows = []
            for line in lines[2:]:  # Skip separator line
                row = [cell.strip() for cell in line.split('|')]
                if len(row) == len(headers):
                    data_rows.append(row)
            
            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Shading Accent 1'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for p in hdr_cells[i].paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add data rows
            for row_data in data_rows:
                cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    cells[i].text = cell_data
                    
                    # Special formatting for evaluation column
                    if i < len(headers) and 'eval' in headers[i].lower():
                        for p in cells[i].paragraphs:
                            for run in p.runs:
                                run.font.bold = True
                                if 'faible' in cell_data.lower():
                                    run.font.color.rgb = RGBColor(0, 128, 0)
                                elif 'moyen' in cell_data.lower():
                                    run.font.color.rgb = RGBColor(255, 165, 0)
                                elif 'élevé' in cell_data.lower():
                                    run.font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph()  # Add spacing after table
            
        except Exception as e:
            logger.warning(f"Table formatting failed: {str(e)}")
            doc.add_paragraph(block)  # Fallback to plain text

    def _add_text_block_to_doc(self, doc: Document, block: str):
        """Add formatted text block to document"""
        lines = block.splitlines()
        if not lines:
            return
        
        first_line = lines[0].strip()
        rest_content = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
        
        # Remove markdown formatting from first line
        clean_first_line = re.sub(r'^#+\s*', '', first_line)
        
        # Determine heading level and add appropriately
        if re.match(r'^#{1,2}\s*\d+\.\s+', first_line) or re.match(r'^\d+\.\s+[A-Z]', clean_first_line):
            # Level 1 heading
            heading = doc.add_heading(clean_first_line, level=1)
            self._format_heading(heading, 1)
        elif re.match(r'^#{2,3}\s*\d+\.\d+\s+', first_line) or re.match(r'^\d+\.\d+\s+', clean_first_line):
            # Level 2 heading
            heading = doc.add_heading(clean_first_line, level=2)
            self._format_heading(heading, 2)
        elif re.match(r'^#{3,4}\s*\d+\.\d+\.\d+\s+', first_line) or re.match(r'^\d+\.\d+\.\d+\s+', clean_first_line):
            # Level 3 heading
            heading = doc.add_heading(clean_first_line, level=3)
            self._format_heading(heading, 3)
        else:
            # Regular paragraph
            doc.add_paragraph(block)
            return
        
        # Add content if present
        if rest_content:
            doc.add_paragraph(rest_content)

    def _format_heading(self, heading, level: int):
        """Format heading with consistent styling"""
        colors = [
            RGBColor(0, 51, 102),   # Level 1
            RGBColor(0, 77, 153),   # Level 2
            RGBColor(0, 102, 204)   # Level 3
        ]
        sizes = [Pt(16), Pt(14), Pt(12)]
        
        for run in heading.runs:
            run.font.bold = True
            run.font.color.rgb = colors[min(level-1, 2)]
            run.font.name = 'Calibri'
            run.font.size = sizes[min(level-1, 2)]

    def _count_tokens(self, text: str, model: str = "gpt-4") -> int:
        """Count tokens in text"""
        try:
            encoding = tiktoken.encoding_for_model(model)
            return len(encoding.encode(text))
        except Exception:
            # Rough estimate if tiktoken fails
            return len(text) // 4


# Factory function
def create_response_generator(
    openai_api_key: Optional[str] = None,
    model: str = "gpt-4-turbo"
) -> ResponseGenerator:
    """Create enhanced response generator with intelligent report generation capabilities"""
    try:
        api_key = openai_api_key or settings.openai.api_key
        generator = ResponseGenerator(
            openai_api_key=api_key,
            model=model
        )
        
        logger.info(f"Enhanced ResponseGenerator created with model {model}")
        logger.info(f"Available report types: proposal, analysis, summary")
        logger.info(f"Available lengths: {list(generator.length_specs.keys())}")
        
        return generator
    except Exception as e:
        logger.error(f"Failed to create response generator: {str(e)}")
        raise AIProcessingException(f"Failed to initialize response generator: {str(e)}")